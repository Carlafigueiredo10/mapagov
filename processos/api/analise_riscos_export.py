"""
Exportacao de Analise de Riscos para Word e PDF

Endpoints:
- GET /api/analise-riscos/<id>/exportar/?formato=docx
- GET /api/analise-riscos/<id>/exportar/?formato=pdf

Arquitetura:
- Snapshot como fonte unica de dados
- SECTIONS define estrutura semantica comum
- Builders geram conteudo por secao
- Renderers (PDF/DOCX) iteram mesma estrutura
"""
import io
import logging
from datetime import datetime
from typing import Any, Dict, List, Optional

from django.http import HttpResponse
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import AllowAny
from rest_framework.response import Response

from processos.models_analise_riscos import AnaliseRiscos
from processos.infra.rate_limiting import rate_limit_user
from processos.api.export_helpers import (
    get_snapshot_para_export,
    verificar_desatualizacao,
    build_snapshot_payload,
    labelize,
    format_value,
    build_bloco_renderizado,
    build_bloco_b_renderizado,
    is_respondido,
    dedup_acoes,
    fmt_text,
    fmt_dash,
    NOTA_CONDICIONAIS,
    LABELS,
)
from processos.domain.helena_analise_riscos.leitura_mgi import (
    gerar_leitura_mgi,
    gerar_leitura_mgi_lista,
    gerar_resumo_mgi,
)

logger = logging.getLogger(__name__)


# =============================================================================
# CONSTANTES
# =============================================================================

def get_orgao_id(request):
    """Extrai orgao_id do usuario (RLS)"""
    import uuid
    if hasattr(request.user, "orgao_id"):
        return request.user.orgao_id
    return uuid.UUID("00000000-0000-0000-0000-000000000000")


# Estrutura semantica unica (PDF e DOCX seguem a mesma ordem)
SECTIONS = [
    {"id": "resumo", "title": "1. Resumo Executivo"},
    {"id": "evidencias", "title": "2. Contexto e Evidencias"},
    {"id": "diagnostico", "title": "3. Diagnostico por Dimensao"},
    {"id": "riscos", "title": "4. Riscos Identificados"},
    {"id": "tratamento", "title": "5. Tratamento dos Riscos"},
    {"id": "leitura_mgi", "title": "6. Enquadramento Institucional e Prioridades Gerenciais (MGI)"},
    {"id": "metadados", "title": "7. Metadados e Rastreabilidade"},
]


# =============================================================================
# BUILDERS - Geram conteudo semantico por secao
# =============================================================================

def build_resumo(data: Dict[str, Any]) -> Dict[str, Any]:
    """
    Constroi conteudo da secao Resumo Executivo.

    Retorna:
        - resumo: totais por nivel
        - nivel_predominante: nivel com mais riscos
        - top_riscos: 3 maiores por score (opcional)
    """
    resumo = data.get("resumo", {})
    riscos = data.get("riscos", [])

    # Nivel predominante
    niveis = ["criticos", "altos", "medios", "baixos"]
    predominante = None
    max_count = 0
    for nivel in niveis:
        count = resumo.get(nivel, 0)
        if count > max_count:
            max_count = count
            predominante = nivel.rstrip("s").upper()  # "criticos" -> "CRITICO"

    # Top 3 riscos por score
    top_riscos = sorted(riscos, key=lambda r: r.get("score_risco") or 0, reverse=True)[:3]

    return {
        "total": resumo.get("total", 0),
        "criticos": resumo.get("criticos", 0),
        "altos": resumo.get("altos", 0),
        "medios": resumo.get("medios", 0),
        "baixos": resumo.get("baixos", 0),
        "nivel_predominante": predominante,
        "top_riscos": [
            {
                "titulo": r.get("titulo", ""),
                "nivel": fmt_text(r.get("nivel_risco"), "Pendente"),
                "score": fmt_dash(r.get("score_risco")),
            }
            for r in top_riscos
        ],
    }


def build_evidencias(data: Dict[str, Any]) -> Dict[str, Any]:
    """
    Constroi conteudo da secao Contexto e Evidencias.

    RETROCOMPATIVEL: Usa helper que faz fallback entre campos novos (v2) e antigos.

    Retorna:
        - bloco_a: campos de identificacao
        - bloco_b: campos de contexto operacional (com fallback)
        - bloco_b_detalhado: estrutura detalhada do Bloco B (v2)
        - nota_condicionais: texto da nota
    """
    contexto = data.get("contexto_estruturado", {})
    bloco_a = contexto.get("bloco_a", {})
    bloco_b = contexto.get("bloco_b", {})

    # Bloco A - campos fixos
    evidencias_a = [
        {"label": "Nome do objeto", "valor": format_value(bloco_a.get("nome_objeto"), required=True)},
        {"label": "Objetivo/Finalidade", "valor": format_value(bloco_a.get("objetivo_finalidade"), required=True)},
        {"label": "Area responsavel", "valor": format_value(bloco_a.get("area_responsavel"), required=True)},
        {"label": "Descricao do escopo", "valor": format_value(bloco_a.get("descricao_escopo"), required=False)},
    ]

    # Bloco B - usa helper com fallback (v2)
    bloco_b_detalhado = build_bloco_b_renderizado(bloco_b)

    # Converte para formato legado (para compatibilidade com templates existentes)
    evidencias_b = [
        {"label": campo["label"], "valor": campo["valor"]}
        for campo in bloco_b_detalhado.get("campos", [])
    ]

    return {
        "bloco_a": evidencias_a,
        "bloco_b": evidencias_b,
        "bloco_b_detalhado": bloco_b_detalhado,  # Estrutura nova (v2)
        "nota_condicionais": NOTA_CONDICIONAIS,
    }


def build_diagnostico(data: Dict[str, Any]) -> List[Dict[str, Any]]:
    """
    Constroi conteudo da secao Diagnostico por Dimensao.

    Retorna lista dos 6 blocos renderizados (usando helper).
    """
    respostas_blocos = data.get("respostas_blocos", {})

    blocos_ids = ["BLOCO_1", "BLOCO_2", "BLOCO_3", "BLOCO_4", "BLOCO_5", "BLOCO_6"]
    diagnostico = []

    for bloco_id in blocos_ids:
        respostas = respostas_blocos.get(bloco_id, {})
        bloco_renderizado = build_bloco_renderizado(bloco_id, respostas)
        diagnostico.append(bloco_renderizado)

    return diagnostico


def build_riscos(data: Dict[str, Any]) -> Dict[str, Any]:
    """
    Constroi conteudo da secao Riscos Identificados.

    Retorna:
        - tabela_resumo: lista para tabela compacta
        - detalhes: lista de blocos detalhados por risco
    """
    riscos = data.get("riscos", [])

    # Tabela resumo (titulo + nivel + score)
    # Usa fmt_dash para evitar None no relatorio
    tabela_resumo = [
        {
            "titulo": r.get("titulo", ""),
            "categoria": labelize(r.get("categoria", "")),
            "nivel": fmt_text(r.get("nivel_risco"), "Pendente"),
            "score": fmt_dash(r.get("score_risco")),
        }
        for r in riscos
    ]

    # Detalhes por risco
    # Usa fmt_dash/fmt_text para renderizacao defensiva (nunca None no relatorio)
    detalhes = []
    for risco in riscos:
        detalhe = {
            "titulo": risco.get("titulo", ""),
            "descricao": format_value(risco.get("descricao"), required=False),
            "categoria": labelize(risco.get("categoria", "")),
            "probabilidade": fmt_dash(risco.get("probabilidade")),
            "impacto": fmt_dash(risco.get("impacto")),
            "score": fmt_dash(risco.get("score_risco")),
            "nivel": fmt_text(risco.get("nivel_risco"), "Pendente"),
            "fonte": labelize(risco.get("fonte_sugestao", "")) or "—",
            "bloco_origem": risco.get("bloco_origem", ""),
            "justificativa": format_value(risco.get("justificativa"), required=False),
            "grau_confianca": labelize(risco.get("grau_confianca", "")) if risco.get("grau_confianca") else "",
            "regra_aplicada": risco.get("regra_aplicada", "") or "",
            "perguntas_acionadoras": risco.get("perguntas_acionadoras", []),
            # NOTA: Plano de resposta REMOVIDO da Secao 4 (Sprint 4)
            # As estrategias e acoes estao descritas na Secao 5
        }

        detalhes.append(detalhe)

    return {
        "tabela_resumo": tabela_resumo,
        "detalhes": detalhes,
        "total": len(riscos),
    }


def build_tratamento(data: Dict[str, Any]) -> Dict[str, Any]:
    """
    Constroi conteudo da secao Tratamento dos Riscos.

    Apresenta o plano de resposta por risco, agrupando por estrategia
    conforme Guia MGI (EVITAR, MITIGAR, COMPARTILHAR, ACEITAR).

    Retorna:
        - riscos_com_resposta: lista de riscos com tratamento definido
        - riscos_sem_resposta: lista de riscos sem tratamento definido
        - totais_por_estrategia: contagem por estrategia
        - total_com_tratamento: quantidade de riscos com tratamento definido
        - total_sem_tratamento: quantidade de riscos sem tratamento definido
    """
    riscos = data.get("riscos", [])

    riscos_com_resposta = []
    riscos_sem_resposta = []

    # Contadores por estrategia (4 do Guia MGI)
    totais_por_estrategia = {
        "EVITAR": 0,
        "MITIGAR": 0,
        "COMPARTILHAR": 0,
        "ACEITAR": 0,
    }

    for risco in riscos:
        respostas = risco.get("respostas", [])
        # Usa funcao canonica para determinar status (REGRA UNICA)
        if is_respondido(risco):
            # Risco com tratamento definido
            # Deduplica acoes para evitar repeticao no relatorio
            respostas_unicas = dedup_acoes(respostas)
            tratamentos = []
            for resp in respostas_unicas:
                estrategia = resp.get("estrategia", "")
                if estrategia in totais_por_estrategia:
                    totais_por_estrategia[estrategia] += 1

                tratamentos.append({
                    "estrategia": labelize(estrategia),
                    "estrategia_raw": estrategia,
                    "acao": resp.get("descricao_acao", ""),
                    "responsavel": resp.get("responsavel_nome", ""),
                    "area": resp.get("responsavel_area", ""),
                    "prazo": resp.get("prazo", ""),
                })

            riscos_com_resposta.append({
                "titulo": risco.get("titulo", ""),
                "nivel": fmt_text(risco.get("nivel_risco"), "Pendente"),
                "score": fmt_dash(risco.get("score_risco")),
                "categoria": labelize(risco.get("categoria", "")),
                "tratamentos": tratamentos,
            })
        else:
            # Risco sem tratamento
            riscos_sem_resposta.append({
                "titulo": risco.get("titulo", ""),
                "nivel": fmt_text(risco.get("nivel_risco"), "Pendente"),
                "score": fmt_dash(risco.get("score_risco")),
                "categoria": labelize(risco.get("categoria", "")),
            })

    return {
        "riscos_com_resposta": riscos_com_resposta,
        "riscos_sem_resposta": riscos_sem_resposta,
        "totais_por_estrategia": totais_por_estrategia,
        "total_com_tratamento": len(riscos_com_resposta),
        "total_sem_tratamento": len(riscos_sem_resposta),
        "nota_tratamento": (
            "O tratamento dos riscos segue as 4 estrategias do Guia de Gestao de Riscos do MGI: "
            "Evitar (eliminar a causa), Mitigar (reduzir probabilidade/impacto), "
            "Compartilhar (transferir a terceiros) e Aceitar (assumir sem acao)."
        ),
        "nota_tratamento_definido": (
            "Tratamento definido NAO significa risco eliminado. "
            "O nivel do risco permanece inalterado ate que as acoes sejam implementadas e avaliadas."
        ),
    }


def build_leitura_mgi(data: Dict[str, Any]) -> Dict[str, Any]:
    """
    Constroi conteudo da secao Leitura MGI.

    Esta secao apresenta a leitura institucional dos riscos conforme
    o Guia de Gestao de Riscos do MGI, sem alterar a classificacao original.

    Retorna:
        - resumo: totais por categoria/nivel MGI
        - riscos_fora_apetite: lista de riscos fora do apetite
        - riscos_integridade: lista de riscos de integridade
        - riscos_com_leitura: lista completa com leitura MGI
    """
    riscos = data.get("riscos", [])
    tipo_origem = data.get("tipo_origem", "")

    # Gerar leitura MGI para cada risco
    riscos_com_leitura = gerar_leitura_mgi_lista(riscos, tipo_origem)

    # Gerar resumo
    resumo = gerar_resumo_mgi(riscos_com_leitura)

    # Identificar riscos fora do apetite SEM resposta definida
    # Usa funcao canonica is_respondido() para consistencia com Secao 5
    riscos_sem_resposta = [
        r for r in riscos_com_leitura
        if (r.get("leitura_mgi") or {}).get("fora_do_apetite", False)
        and not is_respondido(r)
    ]

    return {
        "resumo": resumo,
        "riscos_com_leitura": riscos_com_leitura,
        "riscos_sem_resposta": riscos_sem_resposta,
        "qtd_sem_resposta": len(riscos_sem_resposta),
        "nota_explicativa": (
            "Esta secao apresenta uma camada de leitura institucional dos riscos "
            "conforme o Guia de Gestao de Riscos do MGI. "
            "As classificacoes abaixo NAO substituem a analise original do sistema. "
            "Servem para priorizacao gerencial e reporte institucional."
        ),
        "nota_revisabilidade": (
            "A classificacao MGI, incluindo o status de integridade, reflete a situacao "
            "no momento da analise e pode ser revista caso as causas do risco sejam "
            "tratadas ou eliminadas."
        ),
        "nota_consolidacao": (
            "Alguns riscos apresentam causas ou efeitos comuns e podem ser tratados "
            "de forma integrada, a criterio da gestao."
        ),
        "frase_integridade": (
            "Apetite institucional ZERO. Requerem definicao obrigatoria de resposta."
        ),
        "frase_outros_fora_apetite": (
            "Nivel ALTO ou CRITICO. Prioridade para deliberacao gerencial."
        ),
        "frase_sem_resposta": (
            "Os riscos listados a seguir encontram-se fora do apetite institucional "
            "e ainda nao possuem resposta formal definida."
        ),
    }


def build_metadados(data: Dict[str, Any], snap, desatualizado: bool) -> Dict[str, Any]:
    """
    Constroi conteudo da secao Metadados e Rastreabilidade.

    Se snap=None, indica export stateless (modo teste, sem persistencia).
    """
    # Modo stateless (anônimo) - sem snapshot
    if snap is None:
        return {
            "analise_id": data.get("analise_id", ""),
            "versao_sistema": data.get("versao_sistema", ""),
            "schema_version": data.get("schema_version", ""),
            "fonte_estado_em": data.get("fonte_estado_em", ""),
            "snapshot_versao": "-",
            "snapshot_motivo": "STATELESS (teste)",
            "snapshot_criado_em": None,
            "export_gerado_em": datetime.now().strftime("%d/%m/%Y %H:%M"),
            "desatualizado": False,
            "aviso_desatualizacao": "Export gerado sem snapshot (modo teste publico - sem persistencia).",
        }

    return {
        "analise_id": data.get("analise_id", ""),
        "versao_sistema": data.get("versao_sistema", ""),
        "schema_version": data.get("schema_version", ""),
        "fonte_estado_em": data.get("fonte_estado_em", ""),
        "snapshot_versao": snap.versao,
        "snapshot_motivo": snap.motivo_snapshot,
        "snapshot_criado_em": snap.criado_em.strftime("%d/%m/%Y %H:%M") if snap.criado_em else None,
        "export_gerado_em": datetime.now().strftime("%d/%m/%Y %H:%M"),
        "desatualizado": desatualizado,
        "aviso_desatualizacao": "Ha alteracoes apos o snapshot; este export pode estar desatualizado." if desatualizado else None,
    }


# =============================================================================
# RENDERER PDF (reportlab) - DS Gov "good enough"
# =============================================================================

from processos.export.styles import CORES, CORES_NIVEL, CORES_NIVEL_FUNDO, TIPO, ESPACO

# Cores DS Gov (lazy load para evitar import circular)
COR_AZUL_GOV = None
COR_AZUL_ESCURO = None
COR_CINZA = None
COR_CINZA_CLARO = None

CORES_NIVEL_PDF = None


def _init_pdf_colors():
    """Inicializa cores do PDF (lazy load a partir de styles)."""
    global COR_AZUL_GOV, COR_AZUL_ESCURO, COR_CINZA, COR_CINZA_CLARO, CORES_NIVEL_PDF
    from reportlab.lib import colors

    COR_AZUL_GOV = colors.HexColor(CORES["azul_primario"])
    COR_AZUL_ESCURO = colors.HexColor(CORES["azul_escuro"])
    COR_CINZA = colors.HexColor(CORES["cinza_texto"])
    COR_CINZA_CLARO = colors.HexColor(CORES["cinza_linha"])

    CORES_NIVEL_PDF = {
        nivel: colors.HexColor(hex_cor)
        for nivel, hex_cor in CORES_NIVEL.items()
    }


def _draw_header_footer(canvas, doc, metadados: Dict[str, Any]):
    """Desenha header e footer institucional em todas as paginas."""
    from reportlab.lib.units import cm

    _init_pdf_colors()

    width, height = doc.pagesize

    # === HEADER: Faixa azul no topo ===
    canvas.saveState()

    # Faixa azul
    canvas.setFillColor(COR_AZUL_GOV)
    canvas.rect(0, height - 1.2*cm, width, 1.2*cm, fill=1, stroke=0)

    # Texto do header
    canvas.setFillColor('white')
    canvas.setFont('Helvetica-Bold', 10)
    canvas.drawString(2*cm, height - 0.8*cm, "ANALISE DE RISCOS - HELENA")

    # ID e data no header (lado direito)
    canvas.setFont('Helvetica', 8)
    header_right = f"ID: {metadados.get('analise_id', '')[:8]}... | {metadados.get('export_gerado_em', '')}"
    canvas.drawRightString(width - 2*cm, height - 0.8*cm, header_right)

    canvas.restoreState()

    # === FOOTER: Metadados e paginacao ===
    canvas.saveState()

    # Linha separadora
    canvas.setStrokeColor(COR_CINZA_CLARO)
    canvas.setLineWidth(0.5)
    canvas.line(2*cm, 1.5*cm, width - 2*cm, 1.5*cm)

    # Texto esquerdo: snapshot info
    canvas.setFillColor(COR_CINZA)
    canvas.setFont('Helvetica', 7)
    footer_left = f"Snapshot v{metadados.get('snapshot_versao', '?')} | {metadados.get('versao_sistema', '')} | Schema {metadados.get('schema_version', '')}"
    canvas.drawString(2*cm, 1*cm, footer_left)

    # Texto direito: paginacao
    page_num = canvas.getPageNumber()
    canvas.drawRightString(width - 2*cm, 1*cm, f"Pagina {page_num}")

    canvas.restoreState()


def render_pdf(data: Dict[str, Any], snap, desatualizado: bool) -> io.BytesIO:
    """Renderiza PDF usando SECTIONS e builders com visual DS Gov."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        BaseDocTemplate, PageTemplate, Frame,
        Paragraph, Spacer, Table, TableStyle, KeepTogether
    )

    _init_pdf_colors()

    buffer = io.BytesIO()

    # Metadados para header/footer
    metadados = build_metadados(data, snap, desatualizado)

    # Documento com template customizado
    doc = BaseDocTemplate(
        buffer,
        pagesize=A4,
        topMargin=2.5*cm,  # Espaco para header
        bottomMargin=2*cm,  # Espaco para footer
        leftMargin=2*cm,
        rightMargin=2*cm,
    )

    # Frame principal
    frame = Frame(
        doc.leftMargin,
        doc.bottomMargin,
        doc.width,
        doc.height,
        id='main'
    )

    # Template com header/footer
    def on_page(canvas, doc):
        _draw_header_footer(canvas, doc, metadados)

    template = PageTemplate(id='main', frames=frame, onPage=on_page)
    doc.addPageTemplates([template])

    styles = getSampleStyleSheet()

    # Estilos customizados - DS Gov
    styles.add(ParagraphStyle(
        name='GovTitle',
        parent=styles['Heading1'],
        alignment=1,
        fontSize=18,
        textColor=COR_AZUL_ESCURO,
        spaceBefore=0,
        spaceAfter=4,
    ))
    styles.add(ParagraphStyle(
        name='GovSubtitle',
        parent=styles['Normal'],
        alignment=1,
        fontSize=10,
        textColor=COR_CINZA,
        spaceAfter=16,
    ))
    styles.add(ParagraphStyle(
        name='SectionTitle',
        parent=styles['Heading2'],
        fontSize=13,
        textColor=COR_AZUL_GOV,
        spaceBefore=20,
        spaceAfter=10,
        borderPadding=(0, 0, 4, 0),
        borderWidth=0,
        borderColor=COR_AZUL_GOV,
    ))
    styles.add(ParagraphStyle(
        name='SubSectionTitle',
        parent=styles['Heading3'],
        fontSize=10,
        textColor=COR_AZUL_ESCURO,
        spaceBefore=12,
        spaceAfter=6,
    ))
    styles.add(ParagraphStyle(
        name='GovBody',
        parent=styles['Normal'],
        fontSize=9,
        spaceAfter=4,
        leading=12,
    ))
    styles.add(ParagraphStyle(
        name='SmallGray',
        parent=styles['Normal'],
        fontSize=8,
        textColor=COR_CINZA,
        spaceAfter=2,
    ))
    styles.add(ParagraphStyle(
        name='CardTitle',
        parent=styles['Normal'],
        fontSize=10,
        textColor=COR_AZUL_ESCURO,
        spaceBefore=0,
        spaceAfter=2,
    ))
    styles.add(ParagraphStyle(
        name='Warning',
        parent=styles['Normal'],
        fontSize=TIPO["texto"],
        textColor=colors.HexColor(CORES["vermelho_alerta"]),
        backColor=colors.HexColor(CORES["vermelho_fundo"]),
        borderPadding=ESPACO["sm"],
        spaceBefore=ESPACO["sm"],
        spaceAfter=ESPACO["sm"],
    ))

    elements = []

    # === TITULO DO DOCUMENTO (primeira pagina) ===
    elements.append(Paragraph('MINISTERIO DA GESTAO E DA INOVACAO', styles['GovTitle']))
    elements.append(Paragraph('EM SERVICOS PUBLICOS', styles['GovTitle']))
    elements.append(Spacer(1, 8))
    elements.append(Paragraph('Relatorio de Analise de Riscos', styles['GovSubtitle']))

    # Aviso de desatualizacao (se aplicavel) - logo no inicio
    if metadados['desatualizado'] and metadados['aviso_desatualizacao']:
        elements.append(Paragraph(f"ATENCAO: {metadados['aviso_desatualizacao']}", styles['Warning']))

    elements.append(Spacer(1, 10))

    # Build conteudo por secao
    resumo = build_resumo(data)
    evidencias = build_evidencias(data)
    diagnostico = build_diagnostico(data)
    riscos = build_riscos(data)
    # metadados ja foi construido acima

    # === SECAO 1: RESUMO EXECUTIVO ===
    elements.append(Paragraph(SECTIONS[0]["title"], styles['SectionTitle']))

    # Tabela de distribuicao (visual)
    dist_data = [
        ['Total', 'Criticos', 'Altos', 'Medios', 'Baixos'],
        [
            str(resumo['total']),
            str(resumo['criticos']),
            str(resumo['altos']),
            str(resumo['medios']),
            str(resumo['baixos']),
        ]
    ]
    dist_table = Table(dist_data, colWidths=[3*cm, 2.5*cm, 2.5*cm, 2.5*cm, 2.5*cm])
    dist_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), COR_AZUL_GOV),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('GRID', (0, 0), (-1, -1), 0.5, COR_CINZA_CLARO),
        ('BACKGROUND', (1, 1), (1, 1), colors.HexColor(CORES_NIVEL_FUNDO["CRITICO"])),
        ('BACKGROUND', (2, 1), (2, 1), colors.HexColor(CORES_NIVEL_FUNDO["ALTO"])),
        ('BACKGROUND', (3, 1), (3, 1), colors.HexColor(CORES_NIVEL_FUNDO["MEDIO"])),
        ('BACKGROUND', (4, 1), (4, 1), colors.HexColor(CORES_NIVEL_FUNDO["BAIXO"])),
        ('PADDING', (0, 0), (-1, -1), 8),
    ]))
    elements.append(dist_table)
    elements.append(Spacer(1, 12))

    if resumo['top_riscos']:
        elements.append(Paragraph('<b>Principais riscos (por score):</b>', styles['GovBody']))
        for i, r in enumerate(resumo['top_riscos'], 1):
            cor = CORES_NIVEL_PDF.get(r['nivel'], colors.black)
            elements.append(Paragraph(
                f"{i}. {r['titulo']} - <font color='{cor.hexval()}'><b>{r['nivel']}</b></font> (score: {r['score']})",
                styles['GovBody']
            ))

    elements.append(Spacer(1, 10))

    # === SECAO 2: CONTEXTO E EVIDENCIAS ===
    elements.append(Paragraph(SECTIONS[1]["title"], styles['SectionTitle']))

    # Bloco A
    elements.append(Paragraph('<b>Identificacao do Objeto</b>', styles['SubSectionTitle']))
    for item in evidencias['bloco_a']:
        if item['valor']:
            elements.append(Paragraph(f"<b>{item['label']}:</b> {item['valor']}", styles['GovBody']))

    elements.append(Spacer(1, 6))

    # Bloco B
    elements.append(Paragraph('<b>Contexto Operacional</b>', styles['SubSectionTitle']))
    for item in evidencias['bloco_b']:
        elements.append(Paragraph(f"<b>{item['label']}:</b> {item['valor']}", styles['GovBody']))

    elements.append(Spacer(1, 10))

    # === SECAO 3: DIAGNOSTICO POR DIMENSAO ===
    elements.append(Paragraph(SECTIONS[2]["title"], styles['SectionTitle']))
    elements.append(Paragraph(f"<i>{evidencias['nota_condicionais']}</i>", styles['SmallGray']))
    elements.append(Spacer(1, 6))

    for bloco in diagnostico:
        elements.append(Paragraph(f"<b>{bloco['titulo']}</b>", styles['SubSectionTitle']))

        if not bloco['perguntas']:
            elements.append(Paragraph("<i>Nenhuma resposta registrada</i>", styles['SmallGray']))
        else:
            for p in bloco['perguntas']:
                elements.append(Paragraph(f"• {p['texto']}: <b>{p['valor_label']}</b>", styles['GovBody']))

        elements.append(Spacer(1, 4))

    elements.append(Spacer(1, 10))

    # === SECAO 4: RISCOS IDENTIFICADOS ===
    elements.append(Paragraph(SECTIONS[3]["title"], styles['SectionTitle']))

    if not riscos['detalhes']:
        elements.append(Paragraph("Nenhum risco identificado.", styles['GovBody']))
    else:
        # Tabela resumo compacta (se <= 15 riscos, senao polui)
        if riscos['total'] <= 15:
            table_data = [['Risco', 'Categoria', 'Nivel', 'Score']]
            for r in riscos['tabela_resumo']:
                table_data.append([
                    Paragraph(r['titulo'][:45] + ('...' if len(r['titulo']) > 45 else ''), styles['SmallGray']),
                    r['categoria'],
                    r['nivel'],
                    str(r['score']),
                ])

            t = Table(table_data, colWidths=[8*cm, 3*cm, 2*cm, 1.5*cm])
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), COR_AZUL_GOV),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 0.5, COR_CINZA_CLARO),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('PADDING', (0, 0), (-1, -1), 4),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor(CORES["cinza_fundo"])]),
            ]))
            elements.append(t)
            elements.append(Spacer(1, 16))

        # Detalhes por risco (CARDS)
        elements.append(Paragraph('<b>Detalhamento dos Riscos</b>', styles['SubSectionTitle']))
        elements.append(Spacer(1, 8))

        for i, detalhe in enumerate(riscos['detalhes'], 1):
            cor_nivel = CORES_NIVEL_PDF.get(detalhe['nivel'], colors.black)

            # === CARD DO RISCO (KeepTogether evita quebra no meio) ===
            card_elements = []

            # Titulo do card
            card_elements.append(Paragraph(
                f"<b>{i}. {detalhe['titulo']}</b>",
                styles['CardTitle']
            ))

            # Linha de metricas: Nivel | Score | Categoria | P/I
            metrics_text = f"<font color='{cor_nivel.hexval()}'><b>{detalhe['nivel']}</b></font> | Score: {detalhe['score']} | {detalhe['categoria']} | P:{detalhe['probabilidade']} I:{detalhe['impacto']}"
            card_elements.append(Paragraph(metrics_text, styles['SmallGray']))

            # Fonte
            card_elements.append(Paragraph(f"Fonte: {detalhe['fonte']}", styles['SmallGray']))

            # Descricao
            if detalhe['descricao']:
                card_elements.append(Spacer(1, 4))
                card_elements.append(Paragraph(f"<b>Descricao:</b> {detalhe['descricao']}", styles['GovBody']))

            # Justificativa
            if detalhe['justificativa']:
                card_elements.append(Paragraph(f"<b>Justificativa:</b> {detalhe['justificativa']}", styles['GovBody']))

            # Bloco origem / regra (se inferido)
            if detalhe['bloco_origem']:
                meta = f"Origem: {detalhe['bloco_origem']}"
                if detalhe['regra_aplicada']:
                    meta += f" | Regra: {detalhe['regra_aplicada']}"
                if detalhe['grau_confianca']:
                    meta += f" | Confianca: {detalhe['grau_confianca']}"
                card_elements.append(Paragraph(meta, styles['SmallGray']))

            # NOTA: Plano de Resposta REMOVIDO da Secao 4 (conforme Guia MGI)
            # As estrategias e acoes de tratamento estao na Secao 5.
            card_elements.append(Spacer(1, 2))
            card_elements.append(Paragraph(
                "<i>Veja estrategias e acoes de tratamento na Secao 5.</i>",
                styles['SmallGray']
            ))

            # Linha divisoria (visual de card)
            card_elements.append(Spacer(1, 8))

            # Adiciona card como bloco unico (nao quebra no meio)
            elements.append(KeepTogether(card_elements))

            # Separador entre cards
            elements.append(Spacer(1, 4))

    elements.append(Spacer(1, 16))

    # === SECAO 5: TRATAMENTO DOS RISCOS ===
    tratamento = build_tratamento(data)
    elements.append(Paragraph(SECTIONS[4]["title"], styles['SectionTitle']))

    # Nota explicativa
    elements.append(Paragraph(f"<i>{tratamento['nota_tratamento']}</i>", styles['SmallGray']))
    elements.append(Spacer(1, 10))

    # Tabela resumo por estrategia
    if tratamento['total_com_tratamento'] > 0 or tratamento['total_sem_tratamento'] > 0:
        elements.append(Paragraph('<b>Distribuicao por Estrategia</b>', styles['SubSectionTitle']))
        est_data = [
            ['Estrategia', 'Quantidade'],
            ['Evitar', str(tratamento['totais_por_estrategia']['EVITAR'])],
            ['Mitigar', str(tratamento['totais_por_estrategia']['MITIGAR'])],
            ['Compartilhar', str(tratamento['totais_por_estrategia']['COMPARTILHAR'])],
            ['Aceitar', str(tratamento['totais_por_estrategia']['ACEITAR'])],
        ]
        est_table = Table(est_data, colWidths=[6*cm, 3*cm])
        est_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), COR_AZUL_GOV),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('ALIGN', (1, 0), (1, -1), 'CENTER'),
            ('GRID', (0, 0), (-1, -1), 0.5, COR_CINZA_CLARO),
            ('PADDING', (0, 0), (-1, -1), 6),
        ]))
        elements.append(est_table)
        elements.append(Spacer(1, 12))

        # Resumo geral (linguagem corrigida: "tratamento definido" != "risco resolvido")
        elements.append(Paragraph(
            f"<b>Riscos com tratamento definido:</b> {tratamento['total_com_tratamento']} | "
            f"<b>Riscos sem tratamento definido:</b> {tratamento['total_sem_tratamento']}",
            styles['GovBody']
        ))
        elements.append(Spacer(1, 4))
        elements.append(Paragraph(
            f"<i>{tratamento['nota_tratamento_definido']}</i>",
            styles['SmallGray']
        ))
        elements.append(Spacer(1, 12))

    # Cards de riscos com tratamento definido
    if tratamento['riscos_com_resposta']:
        elements.append(Paragraph('<b>Riscos com Tratamento Definido</b>', styles['SubSectionTitle']))
        elements.append(Spacer(1, 6))

        for risco in tratamento['riscos_com_resposta']:
            card_elements = []
            cor_nivel = CORES_NIVEL_PDF.get(risco['nivel'], colors.black)

            # Titulo do risco
            card_elements.append(Paragraph(
                f"<b>{risco['titulo']}</b>",
                styles['CardTitle']
            ))

            # Linha de situacao (AJUSTE 5: clareza institucional)
            # Determina se esta fora do apetite (ALTO ou CRITICO)
            nivel_upper = (risco['nivel'] or "").upper()
            situacao_apetite = "fora do apetite institucional" if nivel_upper in ("ALTO", "CRITICO") else "dentro do apetite institucional"
            card_elements.append(Paragraph(
                f"Nivel do risco: <font color='{cor_nivel.hexval()}'><b>{risco['nivel']}</b></font> ({situacao_apetite}) | "
                f"Situacao: <b>Tratamento definido</b>",
                styles['SmallGray']
            ))

            # Tratamentos
            for trat in risco['tratamentos']:
                trat_text = f"<b>{trat['estrategia']}:</b> {trat['acao']}"
                if trat['responsavel']:
                    trat_text += f" | Resp: {trat['responsavel']}"
                    if trat['area']:
                        trat_text += f" ({trat['area']})"
                if trat['prazo']:
                    trat_text += f" | Prazo: {trat['prazo']}"
                card_elements.append(Paragraph(trat_text, styles['GovBody']))

            card_elements.append(Spacer(1, 6))
            elements.append(KeepTogether(card_elements))

        elements.append(Spacer(1, 10))

    # Lista de riscos sem tratamento definido (AJUSTE 6: linguagem clara)
    if tratamento['riscos_sem_resposta']:
        elements.append(Paragraph('<b>Riscos sem Tratamento Definido</b>', styles['SubSectionTitle']))
        elements.append(Paragraph(
            "<i>Os riscos abaixo ainda nao possuem plano de resposta definido.</i>",
            styles['SmallGray']
        ))
        elements.append(Spacer(1, 4))

        for risco in tratamento['riscos_sem_resposta']:
            cor_nivel = CORES_NIVEL_PDF.get(risco['nivel'], colors.black)
            # Determina texto do nivel (pode ser "Pendente de classificacao")
            nivel_texto = risco['nivel']
            if nivel_texto == "Pendente":
                nivel_texto = "Pendente de classificacao"
            elements.append(Paragraph(
                f"- {risco['titulo']} | Nivel: <font color='{cor_nivel.hexval()}'>{nivel_texto}</font> | Situacao: Tratamento nao definido",
                styles['GovBody']
            ))

    elements.append(Spacer(1, 16))

    # === SECAO 6: LEITURA SEGUNDO O GUIA DE GR DO MGI ===
    leitura_mgi = build_leitura_mgi(data)
    elements.append(Paragraph(SECTIONS[5]["title"], styles['SectionTitle']))

    # Nota explicativa
    elements.append(Paragraph(f"<i>{leitura_mgi['nota_explicativa']}</i>", styles['SmallGray']))
    elements.append(Spacer(1, 10))

    resumo_mgi = leitura_mgi['resumo']

    # Tabela de distribuicao por categoria MGI
    elements.append(Paragraph('<b>Distribuicao por Categoria MGI</b>', styles['SubSectionTitle']))
    cat_mgi_data = [
        ['Categoria', 'Quantidade'],
        ['Estrategico', str(resumo_mgi['por_categoria_mgi'].get('ESTRATEGICO', 0))],
        ['Operacional', str(resumo_mgi['por_categoria_mgi'].get('OPERACIONAL', 0))],
        ['Integridade', str(resumo_mgi['por_categoria_mgi'].get('INTEGRIDADE', 0))],
    ]
    cat_table = Table(cat_mgi_data, colWidths=[6*cm, 3*cm])
    cat_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), COR_AZUL_GOV),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('ALIGN', (1, 0), (1, -1), 'CENTER'),
        ('GRID', (0, 0), (-1, -1), 0.5, COR_CINZA_CLARO),
        ('PADDING', (0, 0), (-1, -1), 6),
    ]))
    elements.append(cat_table)
    elements.append(Spacer(1, 10))

    # Tabela de distribuicao por nivel MGI
    elements.append(Paragraph('<b>Distribuicao por Nivel MGI</b>', styles['SubSectionTitle']))
    nivel_mgi_data = [
        ['Nivel', 'Quantidade', 'Situacao'],
        ['Pequeno', str(resumo_mgi['por_nivel_mgi'].get('PEQUENO', 0)), 'Dentro do apetite'],
        ['Moderado', str(resumo_mgi['por_nivel_mgi'].get('MODERADO', 0)), 'Apetite institucional'],
        ['Alto', str(resumo_mgi['por_nivel_mgi'].get('ALTO', 0)), 'FORA do apetite'],
        ['Critico', str(resumo_mgi['por_nivel_mgi'].get('CRITICO', 0)), 'FORA do apetite'],
    ]
    nivel_table = Table(nivel_mgi_data, colWidths=[4*cm, 3*cm, 5*cm])
    nivel_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), COR_AZUL_GOV),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('ALIGN', (1, 0), (1, -1), 'CENTER'),
        ('GRID', (0, 0), (-1, -1), 0.5, COR_CINZA_CLARO),
        ('PADDING', (0, 0), (-1, -1), 6),
        # Destacar linhas fora do apetite
        ('BACKGROUND', (2, 3), (2, 3), colors.HexColor('#fef3c7')),  # Alto
        ('BACKGROUND', (2, 4), (2, 4), colors.HexColor('#fee2e2')),  # Critico
    ]))
    elements.append(nivel_table)
    elements.append(Spacer(1, 12))

    # Resumo de riscos fora do apetite
    qtd_fora = resumo_mgi['qtd_fora_apetite']
    qtd_integ = resumo_mgi['qtd_integridade']

    if qtd_fora > 0 or qtd_integ > 0:
        elements.append(Paragraph('<b>Atencao: Riscos que Requerem Decisao Gerencial</b>', styles['SubSectionTitle']))

        if qtd_integ > 0:
            elements.append(Paragraph(
                f"<font color='red'><b>Riscos de Integridade ({qtd_integ}):</b></font> {leitura_mgi['frase_integridade']}",
                styles['GovBody']
            ))
            for ri in resumo_mgi['riscos_integridade']:
                gatilhos = ri.get('gatilhos', [])
                gatilhos_str = f" [gatilhos: {', '.join(gatilhos)}]" if gatilhos else ""
                elements.append(Paragraph(f"  - {ri['titulo']} (score: {ri['score']}){gatilhos_str}", styles['SmallGray']))

        if qtd_fora > qtd_integ:  # Ha outros fora do apetite alem dos de integridade
            qtd_outros = qtd_fora - qtd_integ
            elements.append(Spacer(1, 6))
            elements.append(Paragraph(
                f"<font color='#f97316'><b>Outros riscos fora do apetite ({qtd_outros}):</b></font> {leitura_mgi['frase_outros_fora_apetite']}",
                styles['GovBody']
            ))
            # Usa riscos_com_leitura para ter acesso ao risco completo e verificar tratamento
            for r in leitura_mgi['riscos_com_leitura']:
                lm = r.get('leitura_mgi') or {}
                if lm.get('fora_do_apetite') and not lm.get('is_integridade'):
                    # Verifica situacao do tratamento usando funcao canonica
                    situacao_trat = "Tratamento definido" if is_respondido(r) else "Tratamento nao definido"
                    elements.append(Paragraph(
                        f"  - {r.get('titulo')} ({fmt_text(lm.get('categoria_mgi'), '—')}, {fmt_text(lm.get('nivel_mgi'), '—')}) | <b>{situacao_trat}</b>",
                        styles['SmallGray']
                    ))

    elements.append(Spacer(1, 12))

    # Bloco de riscos sem resposta definida
    if leitura_mgi['qtd_sem_resposta'] > 0:
        elements.append(Paragraph('<b>Riscos Pendentes de Deliberacao</b>', styles['SubSectionTitle']))
        elements.append(Paragraph(f"<i>{leitura_mgi['frase_sem_resposta']}</i>", styles['SmallGray']))
        elements.append(Spacer(1, 4))
        for rs in leitura_mgi['riscos_sem_resposta']:
            lm = rs.get('leitura_mgi', {})
            elements.append(Paragraph(
                f"  - {rs.get('titulo')} ({fmt_text(lm.get('categoria_mgi'), '—')}, {fmt_text(lm.get('nivel_mgi'), '—')}, score: {fmt_dash(rs.get('score_risco'))})",
                styles['SmallGray']
            ))
        elements.append(Spacer(1, 12))

    # Notas institucionais
    elements.append(Paragraph('<b>Notas</b>', styles['SubSectionTitle']))
    elements.append(Paragraph(f"<i>{leitura_mgi['nota_revisabilidade']}</i>", styles['SmallGray']))
    elements.append(Spacer(1, 4))
    elements.append(Paragraph(f"<i>{leitura_mgi['nota_consolidacao']}</i>", styles['SmallGray']))

    elements.append(Spacer(1, 16))

    # === ENCAMINHAMENTO AO GESTOR (AJUSTE 7) ===
    # Texto recomendatorio, sem verbo imperativo, sem obrigacao
    elements.append(Paragraph('<b>Encaminhamento ao gestor do objeto</b>', styles['SubSectionTitle']))
    elements.append(Paragraph(
        "Recomenda-se que a presente analise de riscos seja considerada no plano de acao do objeto, "
        "de modo a subsidiar a definicao e a priorizacao das acoes a serem implementadas.",
        styles['GovBody']
    ))

    elements.append(Spacer(1, 16))

    # === SECAO 7: METADADOS E RASTREABILIDADE ===
    elements.append(Paragraph(SECTIONS[6]["title"], styles['SectionTitle']))

    # Metadados em tabela compacta
    meta_data = [
        ['ID da Analise', metadados['analise_id']],
        ['Versao do Sistema', metadados['versao_sistema']],
        ['Schema', metadados['schema_version']],
        ['Estado capturado em', metadados['fonte_estado_em']],
        ['Snapshot', f"v{metadados['snapshot_versao']} ({metadados['snapshot_motivo']})"],
        ['Export gerado em', metadados['export_gerado_em']],
    ]

    meta_table = Table(meta_data, colWidths=[4.5*cm, 10*cm])
    meta_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 8),
        ('TEXTCOLOR', (0, 0), (-1, -1), COR_CINZA),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('PADDING', (0, 0), (-1, -1), 3),
        ('LINEBELOW', (0, 0), (-1, -2), 0.5, COR_CINZA_CLARO),
    ]))
    elements.append(meta_table)

    # Build PDF
    doc.build(elements)
    buffer.seek(0)

    return buffer


# =============================================================================
# RENDERER DOCX (python-docx)
# =============================================================================

def _hex_to_rgb(hex_color: str):
    """Converte hex (#RRGGBB) para RGBColor do python-docx."""
    from docx.shared import RGBColor
    hex_color = hex_color.lstrip('#')
    return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))


def render_docx(data: Dict[str, Any], snap, desatualizado: bool) -> io.BytesIO:
    """Renderiza DOCX usando SECTIONS e builders - editavel e limpo."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE

    doc = Document()

    # Cores do styles.py convertidas para RGBColor
    COR_AZUL_GOV = _hex_to_rgb(CORES["azul_primario"])
    COR_AZUL_ESCURO = _hex_to_rgb(CORES["azul_escuro"])
    COR_CINZA = _hex_to_rgb(CORES["cinza_texto"])
    COR_VERMELHO = _hex_to_rgb(CORES["vermelho_alerta"])

    CORES_NIVEL_DOCX = {
        nivel: _hex_to_rgb(hex_cor)
        for nivel, hex_cor in CORES_NIVEL.items()
    }

    # === CONFIGURAR ESTILOS GLOBAIS ===
    # Normal
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(10)
    style_normal.paragraph_format.space_after = Pt(6)

    # Heading 1
    style_h1 = doc.styles['Heading 1']
    style_h1.font.name = 'Calibri'
    style_h1.font.size = Pt(14)
    style_h1.font.bold = True
    style_h1.font.color.rgb = COR_AZUL_GOV
    style_h1.paragraph_format.space_before = Pt(18)
    style_h1.paragraph_format.space_after = Pt(10)

    # Heading 2
    style_h2 = doc.styles['Heading 2']
    style_h2.font.name = 'Calibri'
    style_h2.font.size = Pt(12)
    style_h2.font.bold = True
    style_h2.font.color.rgb = COR_AZUL_ESCURO
    style_h2.paragraph_format.space_before = Pt(14)
    style_h2.paragraph_format.space_after = Pt(8)

    # === CABECALHO DO DOCUMENTO ===
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run('MINISTERIO DA GESTAO E DA INOVACAO EM SERVICOS PUBLICOS')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = COR_AZUL_ESCURO

    subheader = doc.add_paragraph()
    subheader.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subheader.add_run('Relatorio de Analise de Riscos')
    run.font.size = Pt(11)
    run.font.color.rgb = COR_CINZA

    # Build metadados
    metadados = build_metadados(data, snap, desatualizado)

    # Aviso de desatualizacao (se aplicavel) logo no inicio
    if metadados['desatualizado'] and metadados['aviso_desatualizacao']:
        p_aviso = doc.add_paragraph()
        run_aviso = p_aviso.add_run(f"ATENCAO: {metadados['aviso_desatualizacao']}")
        run_aviso.bold = True
        run_aviso.font.color.rgb = COR_VERMELHO

    doc.add_paragraph()

    # Build conteudo por secao
    resumo = build_resumo(data)
    evidencias = build_evidencias(data)
    diagnostico = build_diagnostico(data)
    riscos = build_riscos(data)
    # metadados ja foi construido acima

    # === SECAO 1: RESUMO EXECUTIVO ===
    doc.add_heading(SECTIONS[0]["title"], level=1)

    # Tabela de distribuicao
    table_dist = doc.add_table(rows=2, cols=5)
    table_dist.style = 'Table Grid'

    # Header
    headers = ['Total', 'Criticos', 'Altos', 'Medios', 'Baixos']
    for i, h in enumerate(headers):
        cell = table_dist.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Valores
    values = [resumo['total'], resumo['criticos'], resumo['altos'], resumo['medios'], resumo['baixos']]
    for i, v in enumerate(values):
        cell = table_dist.rows[1].cells[i]
        cell.text = str(v)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    if resumo['top_riscos']:
        p2 = doc.add_paragraph()
        p2.add_run('Principais riscos (por score):').bold = True
        for i, r in enumerate(resumo['top_riscos'], 1):
            p3 = doc.add_paragraph(style='List Bullet')
            p3.add_run(f'{r["titulo"]} - ')
            run_nivel = p3.add_run(r['nivel'])
            run_nivel.font.color.rgb = CORES_NIVEL_DOCX.get(r['nivel'], COR_CINZA)
            run_nivel.bold = True
            p3.add_run(f' (score: {r["score"]})')

    doc.add_paragraph()

    # === SECAO 2: CONTEXTO E EVIDENCIAS ===
    h2 = doc.add_heading(SECTIONS[1]["title"], level=1)
    h2.runs[0].font.color.rgb = COR_AZUL_GOV

    # Bloco A
    doc.add_heading('Identificacao do Objeto', level=2)
    for item in evidencias['bloco_a']:
        if item['valor']:
            p = doc.add_paragraph()
            p.add_run(f"{item['label']}: ").bold = True
            p.add_run(item['valor'])

    # Bloco B
    doc.add_heading('Contexto Operacional', level=2)
    for item in evidencias['bloco_b']:
        p = doc.add_paragraph()
        p.add_run(f"{item['label']}: ").bold = True
        p.add_run(item['valor'])

    doc.add_paragraph()

    # === SECAO 3: DIAGNOSTICO POR DIMENSAO ===
    h3 = doc.add_heading(SECTIONS[2]["title"], level=1)
    h3.runs[0].font.color.rgb = COR_AZUL_GOV

    # Nota condicionais
    p_nota = doc.add_paragraph()
    run_nota = p_nota.add_run(evidencias['nota_condicionais'])
    run_nota.italic = True
    run_nota.font.size = Pt(9)
    run_nota.font.color.rgb = COR_CINZA

    for bloco in diagnostico:
        doc.add_heading(bloco['titulo'], level=2)

        if not bloco['perguntas']:
            p = doc.add_paragraph()
            run = p.add_run('Nenhuma resposta registrada')
            run.italic = True
            run.font.color.rgb = COR_CINZA
        else:
            for pergunta in bloco['perguntas']:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"{pergunta['texto']}: ")
                run_resp = p.add_run(pergunta['valor_label'])
                run_resp.bold = True

    doc.add_paragraph()

    # === SECAO 4: RISCOS IDENTIFICADOS ===
    doc.add_heading(SECTIONS[3]["title"], level=1)

    if not riscos['detalhes']:
        doc.add_paragraph('Nenhum risco identificado.')
    else:
        # Tabela resumo compacta (se <= 15 riscos)
        if riscos['total'] <= 15:
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'

            # Header
            headers = ['Risco', 'Categoria', 'Nivel', 'Score']
            for i, h in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = h
                cell.paragraphs[0].runs[0].bold = True

            # Linhas
            for r in riscos['tabela_resumo']:
                row = table.add_row()
                row.cells[0].text = r['titulo'][:50] + ('...' if len(r['titulo']) > 50 else '')
                row.cells[1].text = r['categoria']
                row.cells[2].text = r['nivel']
                row.cells[3].text = str(r['score'])

            doc.add_paragraph()

        # Detalhamento por risco
        doc.add_heading('Detalhamento dos Riscos', level=2)

        for i, detalhe in enumerate(riscos['detalhes'], 1):
            # === CARD DO RISCO ===

            # Titulo
            p_titulo = doc.add_paragraph()
            p_titulo.add_run(f'{i}. {detalhe["titulo"]}').bold = True

            # Linha de metricas
            p_meta = doc.add_paragraph()
            run_nivel = p_meta.add_run(f'{detalhe["nivel"]}')
            run_nivel.bold = True
            run_nivel.font.color.rgb = CORES_NIVEL_DOCX.get(detalhe['nivel'], COR_CINZA)
            run_rest = p_meta.add_run(f' | Score: {detalhe["score"]} | {detalhe["categoria"]} | P:{detalhe["probabilidade"]} I:{detalhe["impacto"]}')
            run_rest.font.size = Pt(9)
            run_rest.font.color.rgb = COR_CINZA

            # Fonte
            p_fonte = doc.add_paragraph()
            run_fonte = p_fonte.add_run(f'Fonte: {detalhe["fonte"]}')
            run_fonte.font.size = Pt(9)
            run_fonte.font.color.rgb = COR_CINZA

            # Descricao
            if detalhe['descricao']:
                p = doc.add_paragraph()
                p.add_run('Descricao: ').bold = True
                p.add_run(detalhe['descricao'])

            # Justificativa
            if detalhe['justificativa']:
                p = doc.add_paragraph()
                p.add_run('Justificativa: ').bold = True
                p.add_run(detalhe['justificativa'])

            # Origem/Regra
            if detalhe['bloco_origem']:
                p_origem = doc.add_paragraph()
                meta_text = f"Origem: {detalhe['bloco_origem']}"
                if detalhe['regra_aplicada']:
                    meta_text += f" | Regra: {detalhe['regra_aplicada']}"
                if detalhe['grau_confianca']:
                    meta_text += f" | Confianca: {detalhe['grau_confianca']}"
                run_origem = p_origem.add_run(meta_text)
                run_origem.font.size = Pt(9)
                run_origem.font.color.rgb = COR_CINZA

            # NOTA: Plano de Resposta REMOVIDO da Secao 4 (conforme Guia MGI)
            # As estrategias e acoes de tratamento estao na Secao 5.
            p_nota_secao5 = doc.add_paragraph()
            run_nota = p_nota_secao5.add_run('Veja estrategias e acoes de tratamento na Secao 5.')
            run_nota.italic = True
            run_nota.font.size = Pt(8)
            run_nota.font.color.rgb = COR_CINZA

            # Separador visual entre riscos
            doc.add_paragraph()

    # === SECAO 5: TRATAMENTO DOS RISCOS ===
    tratamento = build_tratamento(data)
    doc.add_heading(SECTIONS[4]["title"], level=1)

    # Nota explicativa
    p_nota_trat = doc.add_paragraph()
    run_nota_trat = p_nota_trat.add_run(tratamento['nota_tratamento'])
    run_nota_trat.italic = True
    run_nota_trat.font.size = Pt(9)
    run_nota_trat.font.color.rgb = COR_CINZA

    doc.add_paragraph()

    # Tabela resumo por estrategia
    if tratamento['total_com_tratamento'] > 0 or tratamento['total_sem_tratamento'] > 0:
        doc.add_heading('Distribuicao por Estrategia', level=2)
        table_est = doc.add_table(rows=5, cols=2)
        table_est.style = 'Table Grid'

        # Header
        table_est.rows[0].cells[0].text = 'Estrategia'
        table_est.rows[0].cells[1].text = 'Quantidade'
        for cell in table_est.rows[0].cells:
            cell.paragraphs[0].runs[0].bold = True

        # Dados
        est_data = [
            ('Evitar', tratamento['totais_por_estrategia']['EVITAR']),
            ('Mitigar', tratamento['totais_por_estrategia']['MITIGAR']),
            ('Compartilhar', tratamento['totais_por_estrategia']['COMPARTILHAR']),
            ('Aceitar', tratamento['totais_por_estrategia']['ACEITAR']),
        ]
        for i, (est, qtd) in enumerate(est_data, 1):
            table_est.rows[i].cells[0].text = est
            table_est.rows[i].cells[1].text = str(qtd)

        doc.add_paragraph()

        # Resumo geral (linguagem corrigida: "tratamento definido" != "risco resolvido")
        p_resumo_trat = doc.add_paragraph()
        p_resumo_trat.add_run(f"Riscos com tratamento definido: ").bold = True
        p_resumo_trat.add_run(f"{tratamento['total_com_tratamento']} | ")
        p_resumo_trat.add_run(f"Riscos sem tratamento definido: ").bold = True
        p_resumo_trat.add_run(str(tratamento['total_sem_tratamento']))

        # Nota explicativa
        p_nota_def = doc.add_paragraph()
        run_nota_def = p_nota_def.add_run(tratamento['nota_tratamento_definido'])
        run_nota_def.italic = True
        run_nota_def.font.size = Pt(8)
        run_nota_def.font.color.rgb = COR_CINZA

        doc.add_paragraph()

    # Riscos com tratamento definido (AJUSTE 5)
    if tratamento['riscos_com_resposta']:
        doc.add_heading('Riscos com Tratamento Definido', level=2)

        for risco in tratamento['riscos_com_resposta']:
            # Titulo do risco
            p_titulo_risco = doc.add_paragraph()
            p_titulo_risco.add_run(f"{risco['titulo']}").bold = True

            # Linha de situacao (AJUSTE 5: clareza institucional)
            p_situacao = doc.add_paragraph()
            nivel_upper = (risco['nivel'] or "").upper()
            situacao_apetite = "fora do apetite institucional" if nivel_upper in ("ALTO", "CRITICO") else "dentro do apetite institucional"
            p_situacao.add_run(f"Nivel do risco: ")
            run_nivel = p_situacao.add_run(risco['nivel'])
            run_nivel.font.color.rgb = CORES_NIVEL_DOCX.get(risco['nivel'], COR_CINZA)
            run_nivel.bold = True
            p_situacao.add_run(f" ({situacao_apetite}) | Situacao: ")
            p_situacao.add_run("Tratamento definido").bold = True

            # Tratamentos
            for trat in risco['tratamentos']:
                p_trat = doc.add_paragraph(style='List Bullet')
                p_trat.add_run(f"{trat['estrategia']}: ").bold = True
                p_trat.add_run(trat['acao'])
                if trat['responsavel']:
                    p_trat.add_run(f" | Resp: {trat['responsavel']}")
                    if trat['area']:
                        p_trat.add_run(f" ({trat['area']})")
                if trat['prazo']:
                    p_trat.add_run(f" | Prazo: {trat['prazo']}")

            doc.add_paragraph()

    # Riscos sem tratamento definido (AJUSTE 6)
    if tratamento['riscos_sem_resposta']:
        doc.add_heading('Riscos sem Tratamento Definido', level=2)
        p_pend_nota = doc.add_paragraph()
        run_pend_nota = p_pend_nota.add_run('Os riscos abaixo ainda nao possuem plano de resposta definido.')
        run_pend_nota.italic = True
        run_pend_nota.font.size = Pt(9)
        run_pend_nota.font.color.rgb = COR_CINZA

        for risco in tratamento['riscos_sem_resposta']:
            p_pend = doc.add_paragraph(style='List Bullet')
            p_pend.add_run(f"{risco['titulo']} | Nivel: ")
            # Determina texto do nivel
            nivel_texto = risco['nivel']
            if nivel_texto == "Pendente":
                nivel_texto = "Pendente de classificacao"
            run_nivel_pend = p_pend.add_run(nivel_texto)
            run_nivel_pend.font.color.rgb = CORES_NIVEL_DOCX.get(risco['nivel'], COR_CINZA)
            run_nivel_pend.bold = True
            p_pend.add_run(f" | Situacao: Tratamento nao definido")

    doc.add_paragraph()

    # === SECAO 6: LEITURA SEGUNDO O GUIA DE GR DO MGI ===
    leitura_mgi = build_leitura_mgi(data)
    doc.add_heading(SECTIONS[5]["title"], level=1)

    # Nota explicativa
    p_nota = doc.add_paragraph()
    run_nota = p_nota.add_run(leitura_mgi['nota_explicativa'])
    run_nota.italic = True
    run_nota.font.size = Pt(9)
    run_nota.font.color.rgb = COR_CINZA

    doc.add_paragraph()
    resumo_mgi = leitura_mgi['resumo']

    # Tabela de distribuicao por categoria MGI
    doc.add_heading('Distribuicao por Categoria MGI', level=2)
    table_cat = doc.add_table(rows=4, cols=2)
    table_cat.style = 'Table Grid'

    # Header
    table_cat.rows[0].cells[0].text = 'Categoria'
    table_cat.rows[0].cells[1].text = 'Quantidade'
    for cell in table_cat.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True

    # Dados
    cat_data = [
        ('Estrategico', resumo_mgi['por_categoria_mgi'].get('ESTRATEGICO', 0)),
        ('Operacional', resumo_mgi['por_categoria_mgi'].get('OPERACIONAL', 0)),
        ('Integridade', resumo_mgi['por_categoria_mgi'].get('INTEGRIDADE', 0)),
    ]
    for i, (cat, qtd) in enumerate(cat_data, 1):
        table_cat.rows[i].cells[0].text = cat
        table_cat.rows[i].cells[1].text = str(qtd)

    doc.add_paragraph()

    # Tabela de distribuicao por nivel MGI
    doc.add_heading('Distribuicao por Nivel MGI', level=2)
    table_nivel = doc.add_table(rows=5, cols=3)
    table_nivel.style = 'Table Grid'

    # Header
    table_nivel.rows[0].cells[0].text = 'Nivel'
    table_nivel.rows[0].cells[1].text = 'Quantidade'
    table_nivel.rows[0].cells[2].text = 'Situacao'
    for cell in table_nivel.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True

    # Dados
    nivel_data = [
        ('Pequeno', resumo_mgi['por_nivel_mgi'].get('PEQUENO', 0), 'Dentro do apetite'),
        ('Moderado', resumo_mgi['por_nivel_mgi'].get('MODERADO', 0), 'Apetite institucional'),
        ('Alto', resumo_mgi['por_nivel_mgi'].get('ALTO', 0), 'FORA do apetite'),
        ('Critico', resumo_mgi['por_nivel_mgi'].get('CRITICO', 0), 'FORA do apetite'),
    ]
    for i, (nivel, qtd, sit) in enumerate(nivel_data, 1):
        table_nivel.rows[i].cells[0].text = nivel
        table_nivel.rows[i].cells[1].text = str(qtd)
        table_nivel.rows[i].cells[2].text = sit

    doc.add_paragraph()

    # Riscos que requerem decisao gerencial
    qtd_fora = resumo_mgi['qtd_fora_apetite']
    qtd_integ = resumo_mgi['qtd_integridade']

    if qtd_fora > 0 or qtd_integ > 0:
        doc.add_heading('Atencao: Riscos que Requerem Decisao Gerencial', level=2)

        if qtd_integ > 0:
            p_integ = doc.add_paragraph()
            run_integ = p_integ.add_run(f'Riscos de Integridade ({qtd_integ}): ')
            run_integ.bold = True
            run_integ.font.color.rgb = COR_VERMELHO
            p_integ.add_run(leitura_mgi['frase_integridade'])

            for ri in resumo_mgi['riscos_integridade']:
                p_ri = doc.add_paragraph(style='List Bullet')
                gatilhos = ri.get('gatilhos', [])
                gatilhos_str = f" [gatilhos: {', '.join(gatilhos)}]" if gatilhos else ""
                p_ri.add_run(f"{ri['titulo']} (score: {ri['score']}){gatilhos_str}")

        if qtd_fora > qtd_integ:
            qtd_outros = qtd_fora - qtd_integ
            doc.add_paragraph()
            p_outros = doc.add_paragraph()
            run_outros = p_outros.add_run(f'Outros riscos fora do apetite ({qtd_outros}): ')
            run_outros.bold = True
            run_outros.font.color.rgb = CORES_NIVEL_DOCX["ALTO"]  # Laranja
            p_outros.add_run(leitura_mgi['frase_outros_fora_apetite'])

            # Usa riscos_com_leitura para ter acesso ao risco completo e verificar tratamento
            for r in leitura_mgi['riscos_com_leitura']:
                lm = r.get('leitura_mgi') or {}
                if lm.get('fora_do_apetite') and not lm.get('is_integridade'):
                    p_rf = doc.add_paragraph(style='List Bullet')
                    # Verifica situacao do tratamento usando funcao canonica
                    situacao_trat = "Tratamento definido" if is_respondido(r) else "Tratamento nao definido"
                    p_rf.add_run(f"{r.get('titulo')} ({fmt_text(lm.get('categoria_mgi'), '—')}, {fmt_text(lm.get('nivel_mgi'), '—')}) | ")
                    run_sit = p_rf.add_run(situacao_trat)
                    run_sit.bold = True

    doc.add_paragraph()

    # Bloco de riscos sem resposta definida
    if leitura_mgi['qtd_sem_resposta'] > 0:
        doc.add_heading('Riscos Pendentes de Deliberacao', level=2)
        p_sem_resp = doc.add_paragraph()
        run_frase = p_sem_resp.add_run(leitura_mgi['frase_sem_resposta'])
        run_frase.italic = True
        run_frase.font.size = Pt(9)
        run_frase.font.color.rgb = COR_CINZA

        for rs in leitura_mgi['riscos_sem_resposta']:
            lm = rs.get('leitura_mgi', {})
            p_rs = doc.add_paragraph(style='List Bullet')
            p_rs.add_run(f"{rs.get('titulo')} ({fmt_text(lm.get('categoria_mgi'), '—')}, {fmt_text(lm.get('nivel_mgi'), '—')}, score: {fmt_dash(rs.get('score_risco'))})")

    doc.add_paragraph()

    # Notas institucionais
    doc.add_heading('Notas', level=2)
    p_nota1 = doc.add_paragraph()
    run_nota1 = p_nota1.add_run(leitura_mgi['nota_revisabilidade'])
    run_nota1.italic = True
    run_nota1.font.size = Pt(9)
    run_nota1.font.color.rgb = COR_CINZA

    p_nota2 = doc.add_paragraph()
    run_nota2 = p_nota2.add_run(leitura_mgi['nota_consolidacao'])
    run_nota2.italic = True
    run_nota2.font.size = Pt(9)
    run_nota2.font.color.rgb = COR_CINZA

    doc.add_paragraph()

    # === ENCAMINHAMENTO AO GESTOR (AJUSTE 7) ===
    # Texto recomendatorio, sem verbo imperativo, sem obrigacao
    doc.add_heading('Encaminhamento ao gestor do objeto', level=2)
    p_encaminhamento = doc.add_paragraph()
    p_encaminhamento.add_run(
        "Recomenda-se que a presente analise de riscos seja considerada no plano de acao do objeto, "
        "de modo a subsidiar a definicao e a priorizacao das acoes a serem implementadas."
    )

    doc.add_paragraph()

    # === SECAO 7: METADADOS E RASTREABILIDADE ===
    doc.add_heading(SECTIONS[6]["title"], level=1)

    # Metadados em tabela compacta
    table_meta = doc.add_table(rows=6, cols=2)

    meta_rows = [
        ('ID da Analise', metadados['analise_id']),
        ('Versao do Sistema', metadados['versao_sistema']),
        ('Schema', metadados['schema_version']),
        ('Estado capturado em', metadados['fonte_estado_em']),
        ('Snapshot', f"v{metadados['snapshot_versao']} ({metadados['snapshot_motivo']})"),
        ('Export gerado em', metadados['export_gerado_em']),
    ]

    for i, (label, value) in enumerate(meta_rows):
        row = table_meta.rows[i]
        cell_label = row.cells[0]
        cell_value = row.cells[1]

        run_label = cell_label.paragraphs[0].add_run(label)
        run_label.bold = True
        run_label.font.size = Pt(9)
        run_label.font.color.rgb = COR_CINZA

        run_value = cell_value.paragraphs[0].add_run(str(value) if value else '')
        run_value.font.size = Pt(9)
        run_value.font.color.rgb = COR_CINZA

    # Salvar
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer


# =============================================================================
# FUNCOES LEGADAS (mantidas para compatibilidade)
# =============================================================================

def gerar_docx(analise: AnaliseRiscos) -> io.BytesIO:
    """Wrapper legado - usa novo sistema de snapshot."""
    from processos.api.export_helpers import build_snapshot_payload
    data = build_snapshot_payload(analise)
    return render_docx(data, snap=None, desatualizado=False)


def gerar_pdf(analise: AnaliseRiscos) -> io.BytesIO:
    """Wrapper legado - usa novo sistema de snapshot."""
    from processos.api.export_helpers import build_snapshot_payload
    data = build_snapshot_payload(analise)
    return render_pdf(data, snap=None, desatualizado=False)


# =============================================================================
# ENDPOINT
# =============================================================================

@api_view(["GET"])
@permission_classes([AllowAny])  # PROVISORIO: liberado para teste
@rate_limit_user(limit=10, window=60)
def exportar_analise(request, analise_id):
    """
    GET /api/analise-riscos/<id>/exportar/?formato=pdf|docx

    Exporta analise.

    - Autenticado: usa snapshot (persiste se necessario)
    - Anonimo: stateless (zero writes, gera em memoria)
    """
    try:
        orgao_id = get_orgao_id(request)

        analise = AnaliseRiscos.objects.filter(
            id=analise_id, orgao_id=orgao_id
        ).first()

        if not analise:
            return Response({"erro": "Analise nao encontrada"}, status=404)

        formato = request.GET.get('formato', 'pdf').lower()

        # === STATELESS para anonimo (ZERO WRITES) ===
        if not getattr(request.user, "is_authenticated", False):
            # Gera payload em memoria - NAO cria snapshot
            data = build_snapshot_payload(analise)
            snap = None
            desatualizado = False
        else:
            # Fluxo normal: usa/cria snapshot
            snap = get_snapshot_para_export(analise, request.user)
            data = snap.dados_completos
            desatualizado = verificar_desatualizacao(analise, snap)

        if formato == 'docx':
            buffer = render_docx(data, snap, desatualizado)
            response = HttpResponse(
                buffer.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename="analise_riscos_{analise_id}.docx"'
            return response

        elif formato == 'pdf':
            buffer = render_pdf(data, snap, desatualizado)
            response = HttpResponse(
                buffer.getvalue(),
                content_type='application/pdf'
            )
            response['Content-Disposition'] = f'attachment; filename="analise_riscos_{analise_id}.pdf"'
            return response

        else:
            return Response({"erro": f"Formato '{formato}' nao suportado. Use 'pdf' ou 'docx'."}, status=400)

    except ImportError as e:
        logger.exception("Biblioteca de exportacao nao instalada")
        return Response({"erro": f"Biblioteca nao instalada: {str(e)}"}, status=500)
    except Exception as e:
        logger.exception("Erro ao exportar analise")
        return Response({"erro": str(e)}, status=500)
